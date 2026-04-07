"""
Generates: KISE HMIS System Walkthrough.docx
Run with: python3 generate_hmis_doc.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Apply a hex background colour to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top','left','bottom','right']:
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), kwargs.get('val', 'single'))
        border.set(qn('w:sz'), kwargs.get('sz', '4'))
        border.set(qn('w:color'), kwargs.get('color', '000000'))
        tcBorders.append(border)
    tcPr.append(tcBorders)

def heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p

def para(doc, text='', bold=False, italic=False, size=11, color=None, indent=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(11)
    return p

def numbered(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Number')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    for run in p.runs:
        run.font.size = Pt(11)
    return p

def divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_table(doc, headers, rows, col_widths=None,
              header_bg='2E4057', header_fg='FFFFFF',
              alt_bg='F2F6FA'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(*bytes.fromhex(header_fg))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_bg(cell, header_bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = alt_bg if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacing after table
    return table

# ── Document ─────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── Cover Page ───────────────────────────────────────────────────────────────

doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run('KISE HMIS')
tr.bold = True
tr.font.size = Pt(36)
tr.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = subtitle.add_run('Health Management Information System')
sr.font.size = Pt(18)
sr.font.color.rgb = RGBColor(0x04, 0x86, 0x6A)

doc.add_paragraph()

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr2 = tagline.add_run('End-to-End System Walkthrough')
tr2.font.size = Pt(14)
tr2.italic = True
tr2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Kenya Institute of Special Education\n').font.size = Pt(12)
mr = meta.add_run('System Analysis & Business Analysis Team  |  April 2026')
mr.font.size = Pt(11)
mr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ── 1. What is KISE HMIS ─────────────────────────────────────────────────────

heading(doc, '1.  What is KISE HMIS?', level=1, color='2E4057')

para(doc,
     'KISE stands for the Kenya Institute of Special Education — a facility that serves '
     'people living with disabilities, both children and adults. Services offered include '
     'clinical assessments, physiotherapy, occupational therapy, speech therapy, counselling, '
     'educational evaluations, audiology, vision care, assistive technology, and nutrition support.',
     size=11)

para(doc,
     'HMIS stands for Health Management Information System. It is the digital backbone of '
     'KISE — recording every client interaction from the moment they arrive to the moment '
     'they leave, and everything in between.',
     size=11)

para(doc, 'Before this system, KISE managed client records on paper. Key problems included:', size=11)
bullet(doc, 'Client files were lost or duplicated across visits')
bullet(doc, 'Billing was manual, error-prone, and difficult to audit')
bullet(doc, 'No visibility into how many clients a department served on a given day')
bullet(doc, 'No way to know if a child had been assessed before or what services they received')
bullet(doc, 'Insurance claims could not be tracked systematically')

para(doc, 'KISE HMIS solves all of these problems.', size=11, bold=True)

doc.add_paragraph()

# ── 2. Multi-Branch Architecture ─────────────────────────────────────────────

heading(doc, '2.  Multi-Branch Architecture', level=1, color='2E4057')

para(doc,
     'KISE operates across multiple branches. Each branch operates independently in the system — '
     'staff at one branch can only see clients and data for their own branch. '
     'The Head Office (Super Admin) can view all branches simultaneously.',
     size=11)

add_table(doc,
    headers=['Branch Scope', 'Who Can See It'],
    rows=[
        ['Single Branch', 'Receptionist, Nurse, Intake Officer, Cashier, Service Provider'],
        ['All Branches',  'Super Admin, Admin (branch-level full access), Branch Manager'],
    ],
    col_widths=[7, 10],
)

# ── 3. System Users & Roles ──────────────────────────────────────────────────

heading(doc, '3.  System Users and Their Roles', level=1, color='2E4057')

para(doc,
     'The system enforces strict role-based access. Every user logs in at the admin portal '
     'and sees only the modules their role permits — nothing more, nothing less.',
     size=11)

add_table(doc,
    headers=['Role', 'Who They Are', 'What They Can Do'],
    rows=[
        ['Super Admin',      'IT / System Manager',      'Full access to all branches, all modules, system configuration'],
        ['Admin',            'Branch Administrator',      'Full access within their assigned branch'],
        ['Branch Manager',   'Facility Manager',          'Oversight dashboard, reports, service configuration'],
        ['Receptionist',     'Front Desk Officer',        'Register clients, check in/out, manage appointments'],
        ['Triage Nurse',     'Clinical Screener',         'Conduct triage screening, assign risk flags, route clients'],
        ['Intake Officer',   'Assessment Coordinator',    'Collect biodata, functional screening, recommend services, initiate billing'],
        ['Billing Admin',    'Payment Administrator',     'Review and approve invoices for insurance/sponsor cases'],
        ['Cashier',          'Payments Officer',          'Collect payment, record payment method, issue receipts'],
        ['Service Provider', 'Therapist / Specialist',    'Deliver services, document sessions, schedule follow-ups'],
        ['Queue Manager',    'Flow Coordinator',          'Monitor all queues, call clients, manage no-shows'],
    ],
    col_widths=[4, 4.5, 8.5],
)

doc.add_page_break()

# ── 4. The Client Journey ────────────────────────────────────────────────────

heading(doc, '4.  The Client Journey — From Arrival to Exit', level=1, color='2E4057')

para(doc,
     'Every client moves through a series of numbered stages. The system tracks exactly '
     'which stage a client is at, in real time, across every desk in the facility.',
     size=11)

# Journey flow table
flow_table = doc.add_table(rows=9, cols=3)
flow_table.style = 'Table Grid'
flow_table.alignment = WD_TABLE_ALIGNMENT.CENTER

stage_colors = [
    ('①  RECEPTION',      'Register / Check In',                    '1565C0', 'E3F2FD'),
    ('②  TRIAGE',         'Safety Screening',                       '6A1B9A', 'F3E5F5'),
    ('③  INTAKE',         'Assessment & Service Selection',         '2E7D32', 'E8F5E9'),
    ('④  BILLING',        'Invoice Review (insurance cases)',        'E65100', 'FFF3E0'),
    ('⑤  PAYMENT',        'Collect Payment',                        'BF360C', 'FBE9E7'),
    ('⑥  SERVICE POINT',  'Deliver Service',                        '00695C', 'E0F2F1'),
    ('⑦  CHECK-OUT',      'Close the Visit',                        '37474F', 'ECEFF1'),
    ('↩  RETURNING CLIENT SHORTCUT', 'After Triage → skip Intake, go directly to Payment', '795548', 'EFEBE9'),
    ('↻  FOLLOW-UP',      'Scheduled Appointment → Reception → Triage → Payment → Service', '546E7A', 'ECEFF1'),
]

for i, (stage, desc, fg, bg) in enumerate(stage_colors):
    row = flow_table.rows[i]
    row.cells[0].text = ''
    r0 = row.cells[0].paragraphs[0].add_run(stage)
    r0.bold = True
    r0.font.size = Pt(10)
    r0.font.color.rgb = RGBColor(*bytes.fromhex(fg))
    set_cell_bg(row.cells[0], bg)
    row.cells[0].width = Cm(5)

    row.cells[1].text = ''
    r1 = row.cells[1].paragraphs[0].add_run(desc)
    r1.font.size = Pt(10)
    set_cell_bg(row.cells[1], bg)
    row.cells[1].width = Cm(12)

    row.cells[2].text = ''
    set_cell_bg(row.cells[2], bg)
    row.cells[2].width = Cm(0.5)

doc.add_paragraph()

doc.add_page_break()

# ── Stage 1: Reception ────────────────────────────────────────────────────────

heading(doc, '①  Stage 1: Reception', level=2, color='1565C0')
para(doc, 'Role: Receptionist', bold=True, size=11, color='666666')
para(doc,
     'The client arrives at the reception desk. The receptionist\'s first question: '
     '"Have you been here before?"',
     size=11)

heading(doc, 'New Client Registration', level=3)
para(doc, 'The receptionist captures four essential data fields:', size=11)
bullet(doc, 'Full name of the client or child')
bullet(doc, 'Date of birth — or if unknown, an estimated age in years')
bullet(doc, 'Sex (Male / Female / Intersex)')
bullet(doc, 'If a child (below 17): guardian\'s name and phone number')
bullet(doc, 'If an adult (17 and above): their own phone number')

para(doc,
     'On saving, the system auto-generates a permanent KISE Number (UCI) in the format '
     'KISE/A/000126/2026. This is the client\'s unique identifier for life across all visits '
     'and all branches.',
     size=11)

heading(doc, 'Returning Client', level=3)
para(doc,
     'The receptionist searches the system by name, phone number, or KISE card number. '
     'The existing record is retrieved and identity is confirmed.',
     size=11)

heading(doc, 'Check-In', level=3)
para(doc,
     'Once confirmed, the receptionist clicks Check In. The system records the exact date, '
     'time, and receptionist ID. This timestamp marks the official start of the visit and '
     'is later used to calculate total facility time.',
     size=11)

heading(doc, 'Service Availability Check', level=3)
para(doc,
     'The receptionist verifies whether the requested service is available that day. '
     'Customer care staff update availability every morning before 7:00 AM.',
     size=11)
bullet(doc, 'Services available → direct client to Triage')
bullet(doc, 'Services unavailable → advise client, reschedule if needed, check them out')

para(doc,
     'The system creates a Visit record (e.g. VST-20260406-0028) and sets the current '
     'stage to "Reception."',
     size=11, italic=True)

divider(doc)

# ── Stage 2: Triage ───────────────────────────────────────────────────────────

heading(doc, '②  Stage 2: Triage', level=2, color='6A1B9A')
para(doc, 'Role: Triage Nurse', bold=True, size=11, color='666666')
para(doc,
     'Every client — new or returning — must pass through triage without exception. '
     'This is a mandatory safety gate that cannot be bypassed.',
     size=11)

para(doc, 'The triage clinician captures:', size=11)
bullet(doc, 'Vital signs: temperature, height, weight, blood pressure (systolic/diastolic)')
bullet(doc, 'Physical examination findings')
bullet(doc, 'Onset of condition / when concerns began')
bullet(doc, 'Safety risk assessment (self-harm, harm to others)')
bullet(doc, 'Whether emergency or immediate intervention is required')
bullet(doc, 'Whether a referral out is needed (service not available at KISE)')

para(doc,
     'The HMIS automatically calculates a Risk Score and assigns one of three Risk Flags:',
     size=11)

add_table(doc,
    headers=['Risk Flag', 'Condition', 'System Action'],
    rows=[
        ['🟢  Low / Medium', 'Client is stable and fit to proceed',
         'Route to Intake (new) or Payment (returning therapy client)'],
        ['🟡  Medical Hold',  'Unstable vitals, injury, or immediate medical need',
         'Refer to Dispensary; flag record "Medical Hold"; Assessment Coordinator notified; reschedule once cleared'],
        ['🔴  Crisis',        'Active crisis — self-harm, violence, abuse, safeguarding concern',
         'Crisis Management Protocol activated; Assessment Coordinator and crisis team mobilised immediately; record tagged "Crisis – Protocol Initiated"'],
    ],
    col_widths=[3, 5.5, 8.5],
)

para(doc,
     'All triage findings, risk categories, and actions are permanently recorded in the '
     'HMIS and are instantly visible to all downstream service points.',
     size=11, italic=True)

divider(doc)

# ── Stage 3: Intake ───────────────────────────────────────────────────────────

heading(doc, '③  Stage 3: Intake', level=2, color='2E7D32')
para(doc, 'Role: Intake Officer', bold=True, size=11, color='666666')
para(doc,
     'Intake is where clinical need is mapped to specific services and where billing '
     'originates. The client appears in the intake queue automatically after triage clearance.',
     size=11)

heading(doc, 'Section A — Biodata Verification', level=3)
para(doc, 'The officer confirms and completes any missing information:', size=11)
bullet(doc, 'Full demographic details: nationality, county, ward, primary language')
bullet(doc, 'Disability type and onset (congenital or acquired)')
bullet(doc, 'NCPWD number — critical for government funding eligibility')
bullet(doc, 'Level of functioning (Mild / Moderate / Severe / Profound)')
bullet(doc, 'Support needed: Sign Language, Braille, Mobility assistance, Translation')
bullet(doc, 'Guardian / Next of Kin details')
bullet(doc, 'Socio-economic and contact information')

heading(doc, 'Section B — Referral Information', level=3)
bullet(doc, 'Referral source: Self / School / Hospital / Community Worker / Court / Other')
bullet(doc, 'Type of assessment requested (multi-select)')
bullet(doc, 'Reason for referral (narrative)')
bullet(doc, 'Upload of supporting documents: referral letters, prior medical reports')

heading(doc, 'Section C — Functional Screening', level=3)
para(doc,
     'The intake officer screens the client across 8 functional domains. '
     'This guides which service point the client should be posted to.',
     size=11)

add_table(doc,
    headers=['#', 'Domain', 'What Is Being Assessed'],
    rows=[
        ['1', 'Vision',                          'Are there sight concerns?'],
        ['2', 'Hearing',                          'Are there hearing concerns?'],
        ['3', 'Communication & Speech',           'Are there speech or language concerns?'],
        ['4', 'Mobility & Motor Skills',          'Are there movement or motor concerns?'],
        ['5', 'Learning & Cognition',             'Are there learning difficulties or cognitive challenges?'],
        ['6', 'Behaviour & Social Interaction',   'Are there behavioural or social concerns?'],
        ['7', 'Self-care & Daily Living',         'Can the client manage daily activities independently?'],
        ['8', 'Psychosocial / Emotional',         'Are there emotional health or mental wellbeing concerns?'],
    ],
    col_widths=[1, 5, 11],
)

heading(doc, 'Section D — Medical & Developmental History', level=3)
bullet(doc, 'Known medical conditions or diagnoses (list chronic conditions)')
bullet(doc, 'Previous assessments conducted (list)')
bullet(doc, 'Birth or developmental concerns')
bullet(doc, 'Assistive devices currently in use')
bullet(doc, 'Prior therapy: Speech / OT / PT / Counselling / None')

heading(doc, 'Section E — Educational & Occupational Status', level=3)
bullet(doc, 'Current education level: None / ECD / Primary / Secondary / College / Vocational')
bullet(doc, 'School type: Regular / Special / Integrated')
bullet(doc, 'Academic performance concerns and attendance challenges')
bullet(doc, 'Support services already received')

heading(doc, 'Section F — Family & Environmental Factors', level=3)
bullet(doc, 'Family support available')
bullet(doc, 'Socio-economic or environmental barriers')
bullet(doc, 'Caregiver stress indicators')

heading(doc, 'Section G — Service Recommendations  ★ The Billing Trigger', level=3)
para(doc,
     'Based on all information gathered, the intake officer ticks the services this '
     'client should receive. This is the most operationally critical step in the intake process.',
     size=11)

add_table(doc,
    headers=['Recommended Service', 'Department'],
    rows=[
        ['ASD Assessment',                              'Psychological Services / Educational Assessment'],
        ['Learning Disability (LD) Assessment',         'Educational Assessment'],
        ['Intellectual Disability Assessment',           'Psychological Services'],
        ['Physical & Multiple Disability (PMD) Assessment', 'Physiotherapy / Occupational Therapy'],
        ['Occupational Therapy (OT) Assessment',        'Occupational Therapy'],
        ['Physiotherapy (PT) Assessment',               'Physiotherapy'],
        ['Speech & Language Therapy Assessment',        'Speech & Language Therapy'],
        ['Audiology / Hearing Assessment',              'Audiology'],
        ['Vision Assessment',                           'Vision Services'],
        ['Guidance & Counselling',                      'Guidance & Counseling'],
    ],
    col_widths=[9, 8],
)

para(doc,
     'The moment the intake officer saves the form, the following happens automatically:',
     size=11, bold=True)
numbered(doc, 'A Service Booking is created for each ticked service, linked to the client\'s visit')
numbered(doc, 'The client\'s age determines which services are shown — children see child-priced services; adults see adult-priced services')
numbered(doc, 'Service availability is checked for each recommended service')
numbered(doc, 'An Invoice is automatically generated from the service bookings, with prices and insurance calculations computed by the system')

heading(doc, 'Section H — Payment Eligibility Review', level=3)
para(doc, 'The officer enters the client\'s payment method:', size=11)

add_table(doc,
    headers=['Payment Mode', 'Description'],
    rows=[
        ['Cash',                 'Direct cash payment at the desk'],
        ['M-Pesa',               'Mobile money — M-Pesa receipt number recorded'],
        ['SHA',                  'Social Health Authority — government health insurance scheme'],
        ['NCPWD',                'National Council for Persons with Disabilities fund'],
        ['E-Citizen',            'Government e-payment portal'],
        ['Private Insurance',    'AAR, Jubilee, and other private medical insurers'],
    ],
    col_widths=[4.5, 12.5],
)

para(doc, 'Decision logic:', size=11)
bullet(doc, 'Payment method covers the service → client is signposted to their service point and directed to Payment')
bullet(doc, 'Payment method does not cover → check SHA eligibility')
bullet(doc, '  SHA eligible → refer to SHA Officer for registration', level=1)
bullet(doc, '  Not eligible → advise client; visit closed with "Service Deferred" status', level=1)

divider(doc)

# ── Stage 4: Billing ──────────────────────────────────────────────────────────

heading(doc, '④  Stage 4: Billing (Invoice Review)', level=2, color='E65100')
para(doc, 'Role: Billing Admin', bold=True, size=11, color='666666')
para(doc,
     'For clients paying by cash, M-Pesa, or E-Citizen, the invoice proceeds directly '
     'to the Cashier — no billing admin review is required. '
     'For clients using insurance (SHA, NCPWD, or private insurer), the Billing Admin '
     'reviews the invoice before it reaches the Cashier.',
     size=11)

para(doc, 'The invoice shows:', size=11)
bullet(doc, 'Every service booked for this visit')
bullet(doc, 'Full price per service (base_price)')
bullet(doc, 'Amount the insurer covers (covered_amount)')
bullet(doc, 'Amount the client pays out-of-pocket (client_copay)')
bullet(doc, 'Total invoice amount and total client amount')

para(doc,
     'The Billing Admin approves the invoice. Every approval is timestamped with the '
     'staff ID. All invoice changes are permanently logged and cannot be deleted.',
     size=11, italic=True)

divider(doc)

# ── Stage 5: Payment ──────────────────────────────────────────────────────────

heading(doc, '⑤  Stage 5: Payment', level=2, color='BF360C')
para(doc, 'Role: Cashier / SHA Officer', bold=True, size=11, color='666666')
para(doc,
     'The cashier sees the client\'s approved invoice in their payment queue. '
     'They confirm and record the payment.',
     size=11)

para(doc, 'On payment, the system automatically:', size=11)
bullet(doc, 'Marks the invoice as paid (or partially paid if split-payment)')
bullet(doc, 'Records the receipt number (including M-Pesa transaction codes)')
bullet(doc, 'Updates the visit stage to "Proceed to Service"')
bullet(doc, 'Activates the client\'s service bookings in the relevant department queues')

divider(doc)

# ── Stage 6: Service Point ────────────────────────────────────────────────────

heading(doc, '⑥  Stage 6: Service Point', level=2, color='00695C')
para(doc, 'Role: Service Provider (Therapist / Specialist / Assessor)', bold=True, size=11, color='666666')
para(doc,
     'Each department has its own live service queue. The service provider sees clients '
     'booked for their session in order, with priority flags visible.',
     size=11)

heading(doc, 'Service Delivery', level=3)
bullet(doc, 'Customer Care signs the client in at the service block, recording exact arrival time')
bullet(doc, 'The provider opens the session and documents the intervention')
bullet(doc, 'Findings, observations, and progress notes are recorded')
bullet(doc, 'For multi-session therapy (e.g. 12-session OT course), each session is tracked individually')
bullet(doc, 'Progress status: Improving / Stable / Regressing / Completed')

heading(doc, 'Session Outcomes', level=3)
bullet(doc, 'Follow-up needed → appointment is scheduled; client notified by SMS')
bullet(doc, 'Service complete → visit marked ready for check-out')
bullet(doc, 'Exit status documented: Active / Completed Therapy / Transferred / Lost to Follow-up / Deceased')

divider(doc)

# ── Stage 7: Check-Out ────────────────────────────────────────────────────────

heading(doc, '⑦  Stage 7: Check-Out', level=2, color='37474F')
para(doc, 'Role: Receptionist', bold=True, size=11, color='666666')
para(doc,
     'The receptionist clicks Check Out on the client\'s visit record.',
     size=11)

para(doc, 'The system:', size=11)
bullet(doc, 'Records the check-out timestamp')
bullet(doc, 'Calculates total time spent in the facility (check-in to check-out)')
bullet(doc, 'Sets the visit status to "Completed"')
bullet(doc, 'Locks all documentation — records cannot be altered without an audit trail')
bullet(doc, 'Total facility time feeds directly into operational efficiency reports')

doc.add_page_break()

# ── 5. Supporting Systems ─────────────────────────────────────────────────────

heading(doc, '5.  Supporting Systems Running in the Background', level=1, color='2E4057')

heading(doc, 'Service Availability Management', level=2)
para(doc,
     'Customer Care staff at each service block update availability every morning '
     'before 7:00 AM. If a therapist is absent or equipment fails:',
     size=11)
numbered(doc, 'The service is marked Unavailable with a reason code (staff absent, equipment fault, system downtime)')
numbered(doc, 'The system automatically lists all clients booked for that service')
numbered(doc, 'A group SMS is sent via the Celcom SMS gateway notifying affected clients')
numbered(doc, 'Reception and booking modules are blocked from accepting new bookings for that service')
numbered(doc, 'All actions are logged with timestamp and staff ID for accountability')

heading(doc, 'Appointment Scheduling', level=2)
para(doc,
     'Clients can be booked in advance. For returning clients arriving for a scheduled '
     'appointment, the system auto-generates the billing entry at check-in. '
     'They proceed Reception → Triage → Payment, bypassing intake entirely.',
     size=11)

heading(doc, 'Live Queue Management', level=2)
add_table(doc,
    headers=['Queue', 'Who Sees It', 'Purpose'],
    rows=[
        ['Triage Queue',     'Triage Nurse',          'Clients waiting for triage screening'],
        ['Intake Queue',     'Intake Officer',         'Clients waiting for intake assessment'],
        ['Payment Queue',    'Cashier',                'Clients waiting to pay'],
        ['Service Queues',   'Service Provider (one per department)', 'Clients waiting for their service'],
    ],
    col_widths=[4, 5.5, 7.5],
)

heading(doc, 'SMS Notification System', level=2)
para(doc, 'The system sends two types of automated SMS messages to clients:', size=11)
bullet(doc, 'Service Confirmation SMS — confirms that a booked service remains available')
bullet(doc, 'Service Disruption SMS — advises clients to reschedule when a service is cancelled')
para(doc,
     'Group messages are used where possible to reduce SMS volume and cost. '
     'All messages are logged with timestamp and staff ID.',
     size=11)

heading(doc, 'Audit Trail', level=2)
para(doc,
     'Every action in the system — who created a record, who edited it, what changed, '
     'and when — is permanently logged. This audit log cannot be deleted or altered. '
     'It supports compliance, accountability, and dispute resolution.',
     size=11)

doc.add_page_break()

# ── 6. Services & Departments ────────────────────────────────────────────────

heading(doc, '6.  Services and Departments', level=1, color='2E4057')

para(doc,
     'Services are organised by department. Each service has three key classification fields:',
     size=11)
bullet(doc, 'Age Group — who the service is for: Child (< 18 yrs), Adult (≥ 18 yrs), or All Ages')
bullet(doc, 'Category — what type of service it is: Assessment, Therapy, Counseling, Consultation, or Assistive Technology')
bullet(doc, 'Service Type — the technical classification used for billing and reporting')

para(doc, 'Current service catalog and pricing:', size=11)

add_table(doc,
    headers=['Service', 'Age Group', 'Category', 'Price (KShs)', 'Department'],
    rows=[
        ['Physiotherapy Initial Assessment', 'All Ages', 'Assessment', '2,000', 'Physiotherapy'],
        ['Physiotherapy Session',             'All Ages', 'Therapy',    '1,500', 'Physiotherapy'],
        ['Children PT',                       'Child',    'Therapy',    '500',   'Physiotherapy'],
        ['Adult PT',                          'Adult',    'Therapy',    '1,000', 'Physiotherapy'],
        ['Children Hydrotherapy',             'Child',    'Therapy',    '500',   'Physiotherapy'],
        ['Adult Hydrotherapy',                'Adult',    'Therapy',    '1,500', 'Physiotherapy'],
        ['OT Assessment',                     'All Ages', 'Assessment', '2,500', 'Occupational Therapy'],
        ['OT Session',                        'All Ages', 'Therapy',    '1,800', 'Occupational Therapy'],
        ['Children OT',                       'Child',    'Therapy',    '500',   'Occupational Therapy'],
        ['Adult OT',                          'Adult',    'Therapy',    '1,000', 'Occupational Therapy'],
        ['Children Fine Motor',               'Child',    'Therapy',    '500',   'Occupational Therapy'],
        ['Sensory Integration',               'Child',    'Therapy',    '500',   'Occupational Therapy'],
        ['Adult Assessment Consultation',     'Adult',    'Consultation','1,000','Occupational Therapy'],
        ['Speech & Language Assessment',      'All Ages', 'Assessment', '3,000', 'Speech & Language Therapy'],
        ['Speech Therapy Session',            'All Ages', 'Therapy',    '2,000', 'Speech & Language Therapy'],
        ['Children Speech Therapy',           'Child',    'Therapy',    '500',   'Speech & Language Therapy'],
        ['Adult Speech Therapy',              'Adult',    'Therapy',    '1,000', 'Speech & Language Therapy'],
        ['Adult Speech Assessment',           'Adult',    'Assessment', '2,000', 'Speech & Language Therapy'],
        ['Psychological Assessment',          'All Ages', 'Assessment', '4,000', 'Psychological Services'],
        ['Counseling Session',                'All Ages', 'Counseling', '1,500', 'Psychological Services'],
        ['Play Therapy',                      'Child',    'Therapy',    '500',   'Psychological Services'],
        ['Educational Assessment',            'All Ages', 'Assessment', '3,500', 'Educational Assessment'],
        ['Auditory for Adults',               'Adult',    'Assessment', '1,000', 'Audiology'],
        ['Ear Molds (per ear)',               'All Ages', 'Assistive Technology', '2,000', 'Audiology'],
        ['Nutrition Review',                  'All Ages', 'Consultation','500',  'Guidance & Counseling'],
    ],
    col_widths=[5.5, 2.5, 3, 2.5, 4.5],
)

para(doc,
     'Note: Each department is responsible for adding services not yet listed. '
     'Department managers can add and manage their own service catalog directly in the HMIS.',
     size=11, italic=True)

doc.add_page_break()

# ── 7. Insurance & Payment Modes ─────────────────────────────────────────────

heading(doc, '7.  Insurance and Payment Modes', level=1, color='2E4057')

para(doc,
     'KISE accepts multiple payment modes. Insurance coverage is pre-configured per service '
     'so that when a client selects an insurer, the system automatically splits the invoice '
     'into what the insurer pays and what the client pays.',
     size=11)

add_table(doc,
    headers=['Insurer / Mode', 'Type', 'Default Coverage', 'Processing Days'],
    rows=[
        ['Social Health Authority (SHA)',  'Government Scheme', '75%', '14'],
        ['NCPWD',                          'Government Scheme', '50%', '21'],
        ['E-Citizen',                      'Government Portal', '—',   '—'],
        ['AAR Insurance Kenya',            'Private',           '90%', '7'],
        ['Jubilee Insurance',              'Private',           '85%', '10'],
        ['Cash / M-Pesa',                  'Self-pay',          '0%',  'Immediate'],
    ],
    col_widths=[5.5, 4, 3.5, 4],
)

para(doc,
     'Coverage percentages shown are defaults. Actual coverage per service is configured '
     'individually in the Service Insurance Prices catalog.',
     size=11, italic=True)

doc.add_page_break()

# ── 8. Reports ────────────────────────────────────────────────────────────────

heading(doc, '8.  Reports and Data Export', level=1, color='2E4057')

para(doc,
     'Every department generates a daily report. Management can also access '
     'consolidated reports across all departments and branches.',
     size=11)

add_table(doc,
    headers=['Department', 'Report Contains', 'Format'],
    rows=[
        ['Reception',       'Daily total clients seen, classified by sex',                                                 'PDF'],
        ['Triage',          'Daily total clients seen, classified by sex and risk category',                               'PDF'],
        ['Intake',          'Daily total clients seen, classified by intake officer and sex',                              'PDF'],
        ['Billing',         'Client name, service offered, payment mode, amount charged',                                  'Excel + NCPWD reporting format'],
        ['Customer Care',   'Daily total clients seen, classified by service, therapist, and sex',                         'PDF'],
        ['Management',      'Revenue by department, insurance claims, demographic breakdown, service delivery trends',      'Dashboard + Export'],
    ],
    col_widths=[3.5, 9, 4.5],
)

doc.add_page_break()

# ── 9. Real-World Example ─────────────────────────────────────────────────────

heading(doc, '9.  A Real Example: Amina, Age 7', level=1, color='2E4057')

para(doc,
     'The following walkthrough illustrates every step of the system using a real scenario.',
     size=11)

steps = [
    ('08:00 AM — Reception',
     'Amina arrives with her mother. No record found — she is a new client. '
     'The receptionist registers her: name, DOB (age 7), Female, guardian\'s name and phone. '
     'The system assigns KISE/A/000126/2026. Check-in recorded at 08:04 AM. '
     'Visit VST-20260406-0028 is created.'),
    ('08:15 AM — Triage',
     'Amina appears in the triage queue. The nurse completes the screening. '
     'Vital signs are normal. No safety concerns. Risk Flag: Low (Green). '
     'Routing decision: New client → Intake.'),
    ('08:35 AM — Intake',
     'Intake officer opens Amina\'s record. NCPWD number is confirmed. '
     'Functional screening flags: Communication & Speech (Yes), Learning & Cognition (Yes), '
     'Behaviour & Social Interaction (Yes). No prior therapy. '
     'Service recommendations ticked: Children Speech Therapy, Educational Assessment, '
     'OT Assessment. Payment method: NCPWD. System confirms NCPWD covers all three services.'),
    ('08:36 AM — System Auto-Action',
     '3 Service Bookings are created. Invoice INV-20260406-0031 is generated:\n'
     '   Children Speech Therapy   KShs 500    (NCPWD: 250  |  Client: 250)\n'
     '   OT Assessment             KShs 2,500  (NCPWD: 1,250  |  Client: 1,250)\n'
     '   Educational Assessment    KShs 3,500  (NCPWD: 1,750  |  Client: 1,750)\n'
     '   Total Client Pays: KShs 3,250'),
    ('09:10 AM — Billing',
     'Billing Admin reviews and approves the NCPWD invoice.'),
    ('09:20 AM — Payment',
     'Cashier collects KShs 3,250 in cash. Receipt issued. '
     'All three service bookings are now live in their department queues.'),
    ('09:35 AM — Service Point (Speech)',
     'Amina receives her Children Speech Therapy session. '
     'Provider documents findings and intervention notes.'),
    ('10:20 AM — Service Point (OT)',
     'Amina receives her OT Assessment. Provider documents observations and recommends '
     '12 follow-up OT sessions.'),
    ('11:30 AM — Follow-up Scheduled',
     'Follow-up appointment created for next Thursday. '
     'Amina\'s mother receives an SMS confirmation.'),
    ('11:35 AM — Check-Out',
     'Receptionist checks Amina out. Total facility time: 3 hours 31 minutes. '
     'Visit VST-20260406-0028 is marked Completed.'),
]

for (time_label, description) in steps:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    r1 = p.add_run(f'{time_label}\n')
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0x04, 0x86, 0x6A)
    r2 = p.add_run(description)
    r2.font.size = Pt(11)

doc.add_page_break()

# ── 10. What Makes This System Smart ─────────────────────────────────────────

heading(doc, '10.  What Makes KISE HMIS an Enterprise System', level=1, color='2E4057')

smart_points = [
    ('Age-Aware Service Selection',
     'A 7-year-old only sees child-priced services. A 40-year-old only sees adult services. '
     'The system resolves age from date of birth or estimated age automatically — '
     'the intake officer never has to filter manually.'),
    ('Automatic Billing from Intake',
     'The moment an intake officer ticks service recommendations, an invoice is created. '
     'No manual price lookup, no calculation errors, no missed services.'),
    ('Multi-Branch Data Isolation',
     'Each branch operates in its own data silo. Staff cannot see records from other '
     'branches. The Head Office sees everything. This is enforced at the database level.'),
    ('Insurance Integration',
     'SHA, NCPWD, E-Citizen, and private insurers all have pre-configured coverage '
     'percentages per service. The invoice split is automatic.'),
    ('Immutable Audit Trail',
     'Every action — creation, edit, approval, payment — is permanently logged with '
     'the user\'s identity and timestamp. Nothing can be erased.'),
    ('Live Service Availability',
     'Service availability is updated daily by Customer Care. Reception cannot book a '
     'client into an unavailable service. Affected clients are notified by SMS automatically.'),
    ('Lifetime Client Identity',
     'A client\'s KISE Number follows them for life across all visits and all branches. '
     'Every assessment, every payment, every therapy session is linked to that one identifier.'),
    ('Role-Based Access Control',
     'Each staff role sees exactly what they need — no more, no less. '
     'A cashier cannot open a triage record. A triage nurse cannot modify an invoice. '
     'This protects data integrity and patient confidentiality.'),
]

for (title_text, description) in smart_points:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r1 = p.add_run(f'{title_text}:  ')
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)
    r2 = p.add_run(description)
    r2.font.size = Pt(11)

doc.add_paragraph()
divider(doc)
doc.add_paragraph()

# Footer note
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer_p.add_run(
    'This document was prepared by the KISE HMIS System Analysis Team  |  April 2026\n'
    'Kenya Institute of Special Education  |  KISE HMIS v2'
)
fr.font.size = Pt(9)
fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
fr.italic = True

# ── Save ─────────────────────────────────────────────────────────────────────

output_path = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    'KISE HMIS — System Walkthrough.docx'
)
doc.save(output_path)
print(f'Saved: {output_path}')
